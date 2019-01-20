import os
import shutil
import sys
from docx import Document
from pprint import pprint
import json
import time


def main():    
    docx_in = 'masteringlinuxsecurityandhardening.docx'
    show_data = {}

    document = open_docx(docx_in)
    show_title = get_show_title(document)
    show_data.update({"show_title" : show_title})
    mkdir(show_title)
    show_data.update({"seasons" : get_season_titles(document)})
    print(f'Extracting file and episode names from seasons ', end="", flush = True)
    
    for season, title in enumerate(show_data["seasons"]):
        current_season = show_data["seasons"][season]
        title = title["season_title"] 
        season_dir = f'{show_title}/Season {str(season+1)}'       
        mkdir(season_dir)
        episode_titles = get_episode_titles(document, season+1)
        filenames = get_file_names(season+1)
        episode_data_merge = merge(filenames, episode_titles, show_title, season + 1, season_dir)
        current_season.update({"episode":episode_data_merge})

    print('[Done]')
    #save_show_data(show_data)
    #pprint(show_data)




def open_docx(docx_in):
    
    try:
        document = Document(docx_in)
        print(f'Opening {docx_in}')
        return document
    except Exception as err:
        sys.exit(f'Error loading docx file: {err}')



def get_show_title(document):
    '''returns the text from the first center formatted paragraph object in 
the docx file the key of "show_title"'''
    
    for para in document.paragraphs:
        if para.paragraph_format.alignment==1:
            data= para.text
            print('Show Title is: {data}')
            return data
        else:
            print('No show title found')


def get_season_titles(document):
    '''returns an array of key value pairs with keys "season_title": <text
from each paragraph object starting with the string "Section">''' 
    
    season_count=0
    season_data = []
    
    for para in document.paragraphs:
        if para.text.startswith("Section"):
            season_title = {"season_title" : (para.text.split(":")[1]).strip()}
            season_data.append(season_title)
            season_count += 1
    return season_data


def get_episode_titles(document, season_num):
    '''returns array of dicts containing "new_name" : <text from each paragaph
object starting with the number of specified season>'''
    
    episode_data = []
    
    for para in document.paragraphs:
        if para.text.startswith(str(season_num)):
            episode_title = {"new_name" : (para.text.split(" ", 1)[1]).strip()}
            episode_data.append(episode_title)
    return episode_data
    

def get_file_names(season):
    
    file_data = []
    
    for filename in os.listdir("."):
        if filename.startswith("video" + str(season)):
            old_name = {"old_name" : filename}
            file_data.append(old_name)
    return file_data


def merge(filenames, episode_titles, sh_title, seas_num, season_dir):
    '''merges two arrays of key value pairs into one array with 2 key value pairs per index'''
    if len(filenames) == len(episode_titles):
        
        merge_data = []
        
        for enum, old_file in enumerate(filenames):
            
            #ep_name = episode_titles[enum]["new_name"]
            ep_num = enum +1 
            new_filename = f"{sh_title} - s{seas_num:02}e{ep_num:02}.mp4"
            merge_data.append(old_file)
            merge_data[enum].update({"new_name" : str(new_filename)})
            try:
                new_file = f'{season_dir}/{new_filename}'
                if not os.path.exists(new_file):
                    shutil.copy(old_file["old_name"], new_file)
            except Exception as e:
                print(e)
            else:
                print(f'Error, {new_file} already exists.')
        print(f'{seas_num}, ', end="", flush=True)
        return merge_data
    else:
        sys.exit('Error matching filenames to episodes')


def save_show_data(show_data):
    '''exports show_data to <show_title><timestamp>.json file'''
    
    timestr = time.strftime("%Y%m%d-%H%M%S")
    title_4_file = show_data["show_title"].replace(" ", "_")
    file_out = f'{title_4_file}{timestr}.json'
    
    with open(file_out, 'w') as fout:
        json.dump(show_data, fout)


def rename_file(file_data):

    pass
    #os.rename(src, dst, , )


def mkdir(dir_name):
    try:
        if not os.path.exists(dir_name):
            os.makedirs(dir_name)
        else:
            response = input(f'{dir_name} already exists, continue?(yes/no)')
            if response == 'y' or 'yes':
                return
            else:
                sys.exit(f'"{dir_name}" dir already exists, exiting with no change')
    except Exception as e:
        print(e)


if __name__ == "__main__":
    main()